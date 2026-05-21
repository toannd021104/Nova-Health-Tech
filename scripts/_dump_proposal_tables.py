"""Dump all tables from the Alibaba proposal docx into searchable text."""
from docx import Document
import sys, os

doc = Document(r"D:\HaiAnh\Move\Downloads\Cloud\Clouds\Ali\test - Copy\AlibabaCloud_SA_proposal_technical_architecture.docx")

print(f"=== {len(doc.tables)} tables ===\n")
for i, t in enumerate(doc.tables):
    rows_text = []
    for row in t.rows:
        cells = [c.text.strip().replace("\n", " | ") for c in row.cells]
        rows_text.append(" || ".join(cells))
    body = "\n".join(rows_text)
    if any(k in body for k in ["PAI", "EAS", "Content Moderation", "DataWorks", "SDDP", "Model Studio"]):
        print(f"--- Table {i} (matches) ---")
        print(body)
        print()
