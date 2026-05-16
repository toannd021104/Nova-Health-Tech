"""Read the Alibaba Cloud SA proposal docx and dump its content."""
from docx import Document
from pathlib import Path

doc_path = Path(__file__).parent.parent / "AlibabaCloud_SA_proposal_technical_architecture.docx"
doc = Document(str(doc_path))

print("=" * 80)
print(f"Document: {doc_path.name}")
print("=" * 80)

# Iterate paragraphs and tables in order
from docx.oxml.ns import qn

body = doc.element.body
for child in body.iterchildren():
    tag = child.tag.split('}')[-1]
    if tag == 'p':
        # Find paragraph object
        for p in doc.paragraphs:
            if p._element is child:
                style = p.style.name if p.style else "Normal"
                text = p.text.strip()
                if text:
                    if "Heading" in style:
                        print(f"\n[{style}] {text}")
                    else:
                        print(text)
                break
    elif tag == 'tbl':
        for tbl in doc.tables:
            if tbl._element is child:
                print(f"\n[TABLE]")
                for row in tbl.rows:
                    cells = [cell.text.strip().replace("\n", " | ") for cell in row.cells]
                    print(" | ".join(cells))
                print()
                break
